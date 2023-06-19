from flask import Flask, request, jsonify, g
from flask_cors import CORS
import os
import pinecone
from asyncpg.exceptions import UniqueViolationError
import psycopg2
import json
import secrets
from dotenv import load_dotenv
from langchain.llms import OpenAI
from langchain.prompts import (
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
    ChatPromptTemplate,
    PromptTemplate,
    MessagesPlaceholder,
)
from langchain.memory import ConversationBufferMemory, ConversationBufferWindowMemory
from langchain.chains import LLMChain, ConversationChain, RetrievalQA, ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from langchain.schema import (
    AIMessage,
    HumanMessage,
    SystemMessage
)
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import TextLoader
from langchain.vectorstores import Chroma
from docx import Document
from io import BytesIO
from flask_uploads import UploadSet, configure_uploads, IMAGES
from db import get_table_data, delete_row, setup_db
from datetime import datetime, timedelta
import time
import re
from typing import List, Tuple
import spacy
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from werkzeug.security import generate_password_hash, check_password_hash
import uuid
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions
from email_sending import send_email
from typing import Optional
import pdfplumber
import mimetypes
from job_search import call_serp_api
from docx.oxml import parse_xml

app = Flask(__name__)
CORS(app)

load_dotenv()

chains = {}

DATABASE = "data.db"
DB_URL = os.environ.get("DB_URL")

openai_api_key = os.environ.get("OPENAI_API_KEY")
pinecone_api_key = os.environ.get("PINECONE_API_KEY")
app.config["JWT_SECRET_KEY"] = os.environ.get("JWT_SECRET_KEY")

pinecone.init(api_key=pinecone_api_key, environment="us-west4-gcp")
index = pinecone.Index("resume-bot")

jwt = JWTManager(app)

# Load the spaCy English model
nlp = spacy.load('en_core_web_sm')

persist_directory = 'vectordb'

SECTION_KEYWORDS = [
    "overview",
    "summary",
    "profile",
    "objective",
    "education",
    "academic background",
    "work experience",
    "professional experience",
    "experience",
    "employment history",
    "job history",
    "career history",
    "skills",
    "technical skills",
    "core competencies",
    "capabilities",
    "areas of expertise",
    "expertise",
    "projects",
    "portfolio",
    "awards",
    "achievements",
    "accolades",
    "honors",
    "publications",
    "research",
    "certifications",
    "credentials",
    "licenses",
    "training",
    "languages",
    "fluency",
    "multilingual",
    "references",
    "referees",
    "testimonials",
    "professional affiliations",
    "memberships",
    "associations",
    "activities",
    "extracurricular activities",
    "volunteer work",
    "community involvement",
    "leadership",
    "hobbies",
    "interests",
    "personal interests",
]

@app.before_first_request
def create_tables():
    setup_db()


def get_db():
    if "db" not in g:
        g.db = psycopg2.connect(DB_URL)
    return g.db

def close_db(e=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()

app.teardown_appcontext(close_db)

@app.route("/register", methods=["POST"])
def register():
    data = request.get_json()
    print(data)

    # Check if the user already exists
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT * FROM users WHERE email=%s", (data["email"],))
    existing_user = cursor.fetchone()

    if existing_user:
        return jsonify({"error": "User already exists"}), 400

    # Create the new user
    hashed_password = generate_password_hash(data["password"], method="sha256")
    cursor.execute("INSERT INTO users (first_name, last_name, email, job_title, location, password) VALUES (%s, %s, %s, %s, %s, %s)", (data['first_name'], data['last_name'], data['email'], data['job_title'], data['location'], hashed_password))
    db.commit()
    cursor.execute("SELECT * FROM users WHERE email=%s", (data["email"],))
    user = cursor.fetchone()

    return jsonify({"id": user[0], "email": user[5], "first_name": user[1], "last_name": user[2], "job_title": user[3], "location": user[4]}), 201

@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    print(data)

    # Check if the user exists
    db = get_db()
    cur = db.cursor()

    cur.execute("SELECT * FROM users WHERE email=%s", (data["email"],))
    user = cur.fetchone()
    cur.close()

    if not user or not check_password_hash(user[6], data["password"]):
        return jsonify({"error": "Invalid username or password"}), 401

    # Create the access token
    access_token = create_access_token(identity=user[3])

    return jsonify({"access_token": access_token, "id": user[0], "email": user[5], "first_name": user[1], "last_name": user[2], "job_title": user[3], "location": user[4]}), 200

@app.route("/forgot-password", methods=["POST"])
def forgot_password():
    data = request.get_json()
    email = data.get("email")
    url = data.get("frontendUrl")

    if not email:
        return jsonify({"error": "Email is required"}), 400

    db = get_db()
    cur = db.cursor()

    cur.execute("SELECT * FROM users WHERE email=%s", (email, ))
    user = cur.fetchone()

    if not user:
        return jsonify({"error": "No user found with the provided email"}), 404

    # Generate a password reset token
    reset_token = secrets.token_hex(16)

    # Save the token in the database with an expiration time
    cur.execute("INSERT INTO password_reset_tokens (user_id, token, expires_at) VALUES (%s, %s, NOW() + INTERVAL '1 hour')",
                     (user[0], reset_token))
    db.commit()
    cur.close()

    # Send the password reset email
    reset_link = f"{url}/reset-password?token={reset_token}"
    email_subject = "Password Reset Request"
    email_body = f"Please click the following link to reset your password: {reset_link}"
    
    send_email(email_subject, email_body, email)  # Customize this according to your email sending function

    return jsonify({"message": "A password reset link has been sent to your email address"}), 200

@app.route("/reset-password", methods=["POST"])
def reset_password():
    data = request.get_json()
    token = data.get("token")
    new_password = data.get("password")
    print(token)
    print(new_password)

    if not token or not new_password:
        return jsonify({"error": "Token and new password are required"}), 400

    db = get_db()
    cur = db.cursor() 

    token_entry = cur.execute("SELECT * FROM password_reset_tokens WHERE token=%s AND expires_at > NOW()", (token, ))
    token_entry = cur.fetchone()

    if not token_entry:
        return jsonify({"error": "Invalid or expired token"}), 400

    # Update the user's password
    hashed_password = generate_password_hash(new_password, method="sha256")
    cur.execute("UPDATE users SET password=%s WHERE id=%s", (hashed_password, token_entry[1]))
    db.commit()

    # Delete the used password reset token
    cur.execute("DELETE FROM password_reset_tokens WHERE token=%s", (token, ))
    db.commit()
    cur.close()

    return jsonify({"message": "Password has been successfully reset"}), 200


@app.route("/update-password", methods=["PUT"])
def update_password():
    data = request.get_json()
    email = data.get("email")
    new_password = data.get("newPassword")

    if not email or not new_password:
        return jsonify({"type" : "error", "message": "Email and new password are required"}), 400

    db = get_db()
    cur = db.cursor()

    # Find the user by email
    cur.execute("SELECT * FROM users WHERE email=%s", (email, ))
    user = cur.fetchone()

    if not user:
        return jsonify({"type" : "error", "message": "User not found"}), 404

    # Update the user's password
    hashed_password = generate_password_hash(new_password, method="sha256")
    cur.execute("UPDATE users SET password=%s WHERE id=%s", (hashed_password, user[0]))
    db.commit()
    cur.close()

    return jsonify({"type": "success", "id": user[0], "email": user[5], "first_name": user[1], "last_name": user[2], "job_title": user[3], "location": user[4]}), 200

@app.route("/update-profile", methods=["PUT"])
def update_profile():
    data = request.get_json()
    email = data.get("email")
    location = data.get("location")
    job_title = data.get("job_title")
    first_name = data.get("first_name")
    last_name = data.get("last_name")

    db = get_db()
    cur = db.cursor()

    print(email)

    # Find the user by email
    cur.execute("SELECT * FROM users WHERE email=%s", (email, ))
    user = cur.fetchone()
    print(user)

    if not user:
        return jsonify({"type" : "error", "message": "User not found"}), 404

    cur.execute("UPDATE users SET job_title=%s, location=%s, first_name=%s, last_name=%s WHERE id=%s", (job_title, location, first_name, last_name, user[0]))
    db.commit()
    cur.close()

    return jsonify({"type": "success", "id": user[0], "email": user[5], "first_name": user[1], "last_name": user[2], "job_title": user[3], "location": user[4]}), 200


@app.route("/profile", methods=["GET"])
@jwt_required()
def profile():
    current_user = get_jwt_identity()
    
    # Fetch user data from the database
    db = get_db()
    cur = db.cursor()

    cur.execute("SELECT * FROM users WHERE username=%s", (current_user, ))

    user_data = cur.fetchone()

    if not user_data:
        return jsonify({"error": "User not found"}), 404

    return jsonify({"username": user_data[1]}), 200


@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    user_id = request.form.get("id")
    user_embeddings_directory = os.path.join(persist_directory, f"user_{user_id}_embeddings")
    os.makedirs(user_embeddings_directory, exist_ok=True)
    docs = []

    def extract_raw_text(resume_buffer, content_type):
        plain_text = ""
        if content_type == "application/pdf":
            with pdfplumber.open(BytesIO(resume_buffer)) as pdf:
                for page in pdf.pages:
                    plain_text += page.extract_text(x_tolerance=1, y_tolerance=1) + "\n"
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            document = Document(BytesIO(resume_buffer))

            # Extracting text from paragraphs
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if "Hyperlink" in run.style.name:
                        for rel in run._r.rels.values():
                            if "mailto" in rel.reltype:
                                plain_text += rel._target + "\n"
                            elif "http" in rel.reltype:
                                plain_text += rel._target + "\n"
                    else:
                        plain_text += run.text + "\n"

            # Extracting text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            plain_text += paragraph.text + "\n"
            
            # Extracting text from headers and footers
            for section in document.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    plain_text += paragraph.text + "\n"
                
                footer = section.footer
                for paragraph in footer.paragraphs:
                    plain_text += paragraph.text + "\n"

        else:
            raise ValueError("Unsupported file type")
        return plain_text.strip()
        
    def split_resume_into_sections(resume_text: str) -> List[Tuple[str, str]]:
        doc = nlp(resume_text)

        sections = []
        section_starts = []
        for i, token in enumerate(doc):
            # Check if the token is followed by a newline character
            if token.text.lower().strip() in (keyword.lower() for keyword in SECTION_KEYWORDS) and (i == len(doc) - 1 or doc[i + 1].text == "\n"):
                section_starts.append(i)

        for idx, start in enumerate(section_starts):
            section_header = doc[start].text

            # Get the section content
            if idx < len(section_starts) - 1:
                end = section_starts[idx + 1]
                section_content = doc[start + 1:end]
            else:
                section_content = doc[start + 1:]

            sections.append((section_header, section_content.text))

        return sections

    def save_resume_to_database(user_id: int, resume_sections: List[Tuple[str, str]]):
        db = get_db()
        cur = db.cursor()
        
        # Delete existing resume data
        cur.execute("DELETE FROM resume WHERE user_id=%s", (user_id, ))
        
        # Insert full resume
        cur.execute("INSERT INTO resume (user_id, section, content) VALUES (%s, %s, %s)",
                         (user_id, "FULL RESUME", plain_text))
        
        # Insert each section as a new row in the "resume" table
        for section, content in resume_sections:
            cur.execute("INSERT INTO resume (user_id, section, content) VALUES (%s, %s, %s)",
                             (user_id, section, content))
            docs.append(section + " " + content)

        db.commit()

        cur.close()


    try:
        resume_file = request.files["file"]
        content_type, _ = mimetypes.guess_type(resume_file.filename)
        resume_buffer = resume_file.read()
    except Exception as e:
        print("Error reading uploaded file:", e)
        return jsonify(success=False, message="Error reading uploaded file."), 500

    try:
        plain_text = extract_raw_text(resume_buffer, content_type)
        print(plain_text)
        resume_sections = split_resume_into_sections(plain_text)

        embeddings = OpenAIEmbeddings(
                openai_api_key=openai_api_key,
                model="text-embedding-ada-002"
        )

        text_splitter = RecursiveCharacterTextSplitter(
            # Set a really small chunk size, just to show.
            chunk_size = 500,
            chunk_overlap  = 20,
            length_function = len,
        )
        docs_chunked = text_splitter.create_documents([plain_text])
        docs_text = [doc.page_content for doc in docs_chunked]
        docs_embed = embeddings.embed_documents(docs_text)

        doc_ids = [str(uuid.uuid4()) for _ in docs_chunked]
        metadatas = [{"type": "resume", "user": user_id, "text": _.page_content} for _ in docs_chunked]

        index.delete(
            filter={
                "type": {"$eq": "resume"},
                "user": user_id
            }
        )

        index.upsert(vectors=zip(doc_ids, docs_embed, metadatas))

        save_resume_to_database(user_id, resume_sections)

        return jsonify(success=True, message="Resume uploaded and parsed successfully.")
    except Exception as e:
        print("Error parsing resume:", e)
        return jsonify(success=False, message="Error parsing resume."), 500

@app.route("/get-linkedIn", methods=["GET"])
def get_linkedIn():
    pass


@app.route("/fetch-experience-data", methods=["GET"])
def fetch_experience_data():
    try:
        db = get_db()
        cur = db.cursor()

        cur.execute("SELECT content FROM resume")

        rows = cur.fetchall()

        return jsonify(resume=rows)
    except Exception as e:
        print("Error fetching resume data:", e)
        return jsonify(success=False, message="Error fetching resume data."), 500


@app.route("/gpt-api-call", methods=["POST"])
def gpt_api_call():
    db = get_db()
    cur = db.cursor()

    form_values = request.json
    query = form_values["query"]
    user_id = form_values["id"]
    first_name = form_values["first_name"]
    print(user_id)

    embeddings = OpenAIEmbeddings(
        openai_api_key=openai_api_key,
        model="text-embedding-ada-002"
    )

    query_embed = embeddings.embed_query(query)
    print(query_embed)

    try:
        
        resume_docs = index.query(
            vector=query_embed,
            filter={
                "type": {"$eq": "resume"},
                "user": f"""{user_id}"""
            },
            top_k=2,
            include_metadata=True,
            namespace=""
        )

        print(resume_docs)

        resume_docs_texts = resume_docs["matches"]

        resume_docs_texts_flat = [doc["metadata"]["text"] for doc in resume_docs_texts]

        resume_docs_str = "\n\n".join(resume_docs_texts_flat)

        print(resume_docs_str)

    except:

        resume_docs_str = ""

    try:
        
        jobs_docs = index.query(
            vector=query_embed,
            filter={
                "type": {"$eq": "jobs"},
                "user": f"""{user_id}"""
            },
            top_k=4,
            include_metadata=True
        )

        jobs_docs_texts = jobs_docs["matches"]

        jobs_docs_texts_flat = [doc["metadata"]["text"] for doc in jobs_docs_texts]

        jobs_docs_str = "\n\n".join(jobs_docs_texts_flat)

        context = f"""\nInformation from {first_name}'s resume:\n {resume_docs_str}""" + f"""\nInformation from {first_name}'s job applications:\n {jobs_docs_str}"""

    except:

        context = f"""\nInformation from {first_name}'s resume:\n {resume_docs_str}"""

    try:

        questions_docs = index.query(
            vector=query_embed,
            filter={
                "type": {"$eq": "questions"},
                "user": f"""{user_id}"""
            },
            top_k=1,
            include_metadata=True
        )

        questions_docs_texts = questions_docs["matches"]
        questions_docs_texts_flat = [doc["metadata"]["text"] for doc in questions_docs_texts]
        questions_docs_str = "\n\n".join(questions_docs_texts_flat)

        context = context + f"""\nInterview questions that have been answered by {first_name}:\n {questions_docs_str}"""

    except:
        
        context = context

    chat = ChatOpenAI(
        model_name="gpt-3.5-turbo",
        openai_api_key=openai_api_key,
        temperature=0,
    )

    template_chat_1 = f"""You are an intelligent career personal assistant and your job is to help {first_name} with all things career related. You may be asked to help improve their resume, help with their job applications, prepare for interviews, answer recruiter questions and more. Below is relevant information you can use to answer any questions {first_name} may have. Do not make anything up.\n""" 

    template_chat = template_chat_1 + f"""{context}\n"""

    final_template = template_chat + """{history}\n User question: {input}\n Assistant:"""


    resume_prompt = PromptTemplate(
        input_variables=['history','input'],
        template=final_template      
        )

    chain_id = form_values.get("chainId")

    chain = chains.get(chain_id)

    chat_id = None

    if not chain:
        chain = ConversationChain(
            llm=chat,
            prompt=resume_prompt,
            verbose=True,
            memory=ConversationBufferWindowMemory(k=2),
        )

        cur.execute("INSERT INTO chat (user_id, chat_name) VALUES (%s, %s) RETURNING id", (user_id, query))
        chat_id = cur.fetchone()[0]

        chains[chat_id] = chain

        db.commit()
    
    try:
        result_endpoint = chain.predict(input=query)
        answer = result_endpoint
        print(answer)

        timestamp = datetime.now().isoformat()
        cur.execute(
        "INSERT INTO messages (type, user_id, chat_id, message, timestamp) VALUES (%s, %s, %s, %s, %s)",
        ("user", user_id, chat_id, query, timestamp)
        )
        cur.execute(
            "INSERT INTO messages (type, user_id, chat_id, message, timestamp) VALUES (%s, %s, %s, %s, %s)",
            ("bot", user_id, chat_id, answer, timestamp)
        )
        db.commit()
        cur.close()

        return jsonify(answer=answer, chainId=chain_id)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error fetching GPT-3.5 API"), 500

@app.route("/get-answer-help", methods=["POST"])
def get_answer_help():
    db = get_db()
    cur = db.cursor()

    form_values = request.json
    query = form_values["query"]
    user_id = form_values["id"]
    question_id = form_values["question_id"]
    question_answer = form_values.get("question_answer")
    job_id = form_values.get("job_id")
    print(user_id)

    embeddings = OpenAIEmbeddings(
        openai_api_key=openai_api_key,
        model="text-embedding-ada-002"
    )

    query_embed = embeddings.embed_query(query)

    resume_docs = index.query(
        vector=query_embed,
        filter={
            "type": {"$eq": "resume"},
            "user": f"""{user_id}"""
        },
        top_k=2,
        include_metadata=True,
        namespace=""
    )

    resume_docs_texts = resume_docs["matches"]

    resume_docs_texts_flat = [doc["metadata"]["text"] for doc in resume_docs_texts]

    resume_docs_str = "\n\n".join(resume_docs_texts_flat)

    try:
        questions_docs = index.query(
            vector=query_embed,
            filter={
                "type": {"$eq": "questions"},
                "user": f"""{user_id}"""
            },
            top_k=2,
            include_metadata=True,
            namespace=""
        )
        
        questions_docs_texts = questions_docs["matches"]
        questions_docs_texts_flat = [doc["metadata"]["text"] for doc in questions_docs_texts]
        questions_docs_str = "\n\n".join(questions_docs_texts_flat)
    except:
        questions_docs_str = ""

    chat = ChatOpenAI(
        model_name="gpt-3.5-turbo",
        openai_api_key=openai_api_key,
        temperature=0,
    )

    if question_answer:

        template_base = """You are an intelligent career bot and your job is to help users improve their resume, speak more effectively about their work experience, and provide career guidance. In this case, the user is trying to prepare answers to interview questions specific to their role.

        Please improve the user's answer based on the information below and provide a recommendation to them on how they can better answer the question. 
        
        Your response must be in JSON format with the following properties: "answer": "<Your answer here>", "recommendation": "<Your recommendation here>". Relevant information is provided below. Do not make anything up.\n""" + f"""User's current answer: {question_answer}\n"""
    else:

        template_base = """You are an intelligent career bot and your job is to help users improve their resume, speak more effectively about their work experience, and provide career guidance. In this case, the user is trying to prepare answers to interview questions specific to their role.

        Please answer the question from the user's perspective and provide a recommendation to the user on what they need to include to improve the answer you provided. Your response must be in JSON format with the following properties: "answer": "<Your answer here>", "recommendation": "<Your recommendation here>". Relevant information is provided below. Do not make anything up.\n"""

    if job_id:

        job_docs = index.query(
            vector=query_embed,
            filter={
                "type": {"$eq": "jobs"},
                "user": f"""{user_id}"""
            },
            top_k=2,
            include_metadata=True,
            namespace=""
        )

        job_docs_texts = job_docs["matches"]
        job_docs_texts_flat = [doc["metadata"]["text"] for doc in job_docs_texts]
        job_docs_str = "\n\n".join(job_docs_texts_flat)

        template_help = template_base + f"""\n\nInformation from the job post the user is applying to:\n {job_docs_str}""" + f"""\n\nInformation from the user's resume:\n {resume_docs_str}"""

    else:

        template_help = template_base + f"""\n\nInformation from the user's resume:\n {resume_docs_str}"""

    template_final = template_help + """\nInterview question: {context}"""

    print(template_final)


    chain = LLMChain(
        llm=chat,
        verbose=True,
        prompt=PromptTemplate.from_template(template_final)
    )
    
    try:
        result_endpoint = chain.run(query)
        answer = json.loads(result_endpoint)
        print(answer)

        cur.execute(
            "UPDATE interview_questions SET recommendation = %s WHERE id = %s",
            (answer["recommendation"], question_id),
        )

        db.commit()  # Commit the changes to the database
        cur.close()

        return jsonify(answer)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error fetching GPT-3.5 API"), 500

@app.route("/improve-answer", methods=["POST"])
def improve_answer():
    db = get_db()
    cur = db.cursor()

    form_values = request.json
    query = form_values["query"]
    user_id = form_values["id"]
    question_id = form_values["question_id"]
    question_answer = form_values.get("question_answer")
    job_id = form_values.get("job_id")
    print(user_id)

    embeddings = OpenAIEmbeddings(
        openai_api_key=openai_api_key,
        model="text-embedding-ada-002"
    )

    query_embed = embeddings.embed_query(query)

    resume_docs = index.query(
        vector=query_embed,
        filter={
            "type": {"$eq": "resume"},
            "user": f"""{user_id}"""
        },
        top_k=2,
        include_metadata=True,
        namespace=""
    )

    resume_docs_texts = resume_docs["matches"]

    resume_docs_texts_flat = [doc["metadata"]["text"] for doc in resume_docs_texts]

    resume_docs_str = "\n\n".join(resume_docs_texts_flat)

    job_docs = index.query(
        vector=query_embed,
        filter={
            "type": {"$eq": "jobs"},
            "job_id": str(job_id),
            "user": f"""{user_id}"""
        },
        top_k=2,
        include_metadata=True,
        namespace=""
    )

    job_docs_texts = job_docs["matches"]
    job_docs_texts_flat = [doc["metadata"]["text"] for doc in job_docs_texts]
    job_docs_str = "\n\n".join(job_docs_texts_flat)

    chat = ChatOpenAI(
        model_name="gpt-3.5-turbo",
        openai_api_key=openai_api_key,
        temperature=0,
    )

    if question_answer:

        template_base = """You are an intelligent career bot and your job is to help users improve their resume, speak more effectively about their work experience, and provide career guidance. In this case, the user is trying to improve their answer to an interview question.

        Please suggest improvements to the user's answer based on the information below. Provide a more improved answer and also recommendations on how to approach similar questions in future. The improved answer should be better worded and include any relevant information from the information provided to you.

        Your response must be in JSON with the following properties: "answer": "" , "recommendation": "". Relevant information is provided below. Do not make anything up.\n""" + f"""Interview question user is answering: {query}\n"""

    template_final = template_base + """\nUser's interview question answer: {context}""" + f"""\nInformation from the users resume:\n {resume_docs_str}""" + f"""\nInformation from the job application:\n {job_docs_str}"""

    chain = LLMChain(
        llm=chat,
        verbose=True,
        prompt=PromptTemplate.from_template(template_final)
    )

    try:
        result_endpoint = chain.run(question_answer)
        improved_answer = json.loads(result_endpoint)
        print(improved_answer)

        db.commit()  # Commit the changes to the database
        cur.close()

        return jsonify(improved_answer)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error fetching GPT-3.5 API"), 500

@app.route("/generate-interview-questions", methods=["POST"])
def generate_interview_questions():
    form_values = request.json
    print(form_values)
    user_id = form_values["user_id"]
    job_id = form_values.get("job_id")
    question_type = form_values["type"]

    chat = ChatOpenAI(
        model_name="gpt-3.5-turbo",
        openai_api_key=openai_api_key,
        temperature=0,
    )

    db = get_db()
    cur = db.cursor()

    cur.execute(f"SELECT * FROM interview_questions WHERE user_id = %s AND job_id = %s", (user_id, job_id))

    existing_questions = cur.fetchall()
    questions_list = [row[3] for row in existing_questions]
    questions_str = ', '.join(f'"{question}"' for question in questions_list)

    if question_type == 'WorkExperience':
        type_string = "that are specifically relevant to their work experience"
        query = "Work Experience"
    
    if question_type == "RoleBased":
        type_string = "that are specifically relevant to the job"
        query = "Work Experience"

    if question_type == "Technical":
        type_string = "that are specifically relevant to their technical capabilities"
        query = "Technical Skills"

    embeddings = OpenAIEmbeddings(
        openai_api_key=openai_api_key,
        model="text-embedding-ada-002"
    )

    query_embed = embeddings.embed_query(query)

    resume_docs = index.query(
        vector=query_embed,
        filter={
            "type": {"$eq": "resume"},
            "user": f"""{user_id}"""
        },
        top_k=2,
        include_metadata=True,
        namespace=""
    )

    resume_docs_texts = resume_docs["matches"]

    resume_docs_texts_flat = [doc["metadata"]["text"] for doc in resume_docs_texts]

    resume_docs_str = "\n\n".join(resume_docs_texts_flat)
 
    if job_id:

        resume = resume_docs_str

        cur.execute(f"SELECT * FROM saved_jobs WHERE id = %s AND user_id = %s", (job_id, user_id,))

        job = cur.fetchall()

        print(job[0][4])

        template_questions_base = f"""Below is a user's resume and the description for a job they are applying to. Based on the job responsibilities and user's qualifications, please generate 3 interview questions {type_string}. The questions should be related to key skills and experiences required for this job. Please return the questions separated by newlines.
        
        Job Description user is applying to:\n {job[0][4]}\n
        """

    else:
        template_questions_base = f"""Below is a user's resume and the description for a job they are applying to. Based on the job responsibilities and user's qualifications, please generate 3 interview questions {type_string}. The questions should be related to key skills and experiences required for this job. Please return the questions separated by newlines."""

        cur.execute(f"SELECT * FROM resume WHERE section = %s AND user_id = %s", ("FULL RESUME",user_id,))
        resume = cur.fetchall()

    if existing_questions:
        exclude_similar = f"""These questions have already been asked: [{questions_str}]"""
    else:
        exclude_similar = ""

    template_final = template_questions_base + """Information from the user's Resume:\n {context} \n""" + exclude_similar

    chain = LLMChain(
        llm=chat,
        verbose=True,
        prompt=PromptTemplate.from_template(template_final)
    )

    chain_id = str(int(time.time()))  # Generate a unique identifier based on the current timestamp

    try:
        questions_response = chain(resume)
        print(questions_response)
        # Convert the questions string into a Python list
        questions_string = questions_response['text']
        questions_list = [re.sub(r'^\d+\.\s*', '', i) for i in questions_string.split('\n')]
        # Save the questions to the interview_questions table
        for question in questions_list:
            cur.execute(
                "INSERT INTO interview_questions (user_id, job_id, question) VALUES (%s, %s, %s)",
                (user_id, job_id, question),
            )

        # Commit the changes to the database
        db.commit()
        cur.close()

        return jsonify(questions=questions_list, chainId=chain_id)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error fetching GPT-3.5 API"), 500


@app.route("/get-interview-questions", methods=["POST"])
def get_interview_questions():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    job_id = data.get("job_id")  # Get job_id if it exists in the request

    if job_id:
        # Filter interview questions by user_id and job_id
        cur.execute("SELECT * FROM interview_questions WHERE user_id = %s AND job_id = %s", (user_id, job_id))
    else:
        # Filter interview questions only by user_id
        cur.execute("SELECT * FROM interview_questions WHERE user_id = %s", (user_id,))

    interview_questions = cur.fetchall()

    questions = [
        {
            "id": question[0],
            "user_id": question[1],
            "question": question[3],
            "answer": question[4],
            "job_id": question[2],
            "recommendation": question[5]
        }
        for question in interview_questions
    ]

    return jsonify(questions=questions)

@app.route("/delete-interview-question", methods=["POST"])
def delete_interview_question():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    question_id = data["question_id"]

    user_embeddings_directory = os.path.join(persist_directory, f"user_{user_id}_embeddings")
    os.makedirs(user_embeddings_directory, exist_ok=True)

    try:

        index.delete(
            ids=[str(question_id)],
            namespace=""
        )

        cur.execute("DELETE FROM interview_questions WHERE id = %s AND user_id = %s", (question_id, user_id))
        db.commit()
        cur.close()
        return jsonify(success=True, message="Interview question deleted successfully")
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error deleting the interview question")


@app.route("/save-answer", methods=["POST"])
def save_answer():

    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    question_id = data["question_id"]
    question = data["question"]
    answer = data["answer"]

    try:
        cur.execute(
            "UPDATE interview_questions SET answer = %s WHERE id = %s",
            (answer, question_id),
        )
        db.commit()
        cur.close()

        embeddings = OpenAIEmbeddings(
            openai_api_key=openai_api_key,
            model="text-embedding-ada-002"
        )

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = 5000,
            chunk_overlap  = 20,
            length_function = len,
        )
        docs_chunked = text_splitter.create_documents([f"Question: {question}, Answer: {answer}"])
        docs_text = [doc.page_content for doc in docs_chunked]
        docs_embed = embeddings.embed_documents(docs_text)
        doc_ids = [f"{question_id}"]

        metadatas = [{"type": "questions", "user": str(user_id), "text": _.page_content} for _ in docs_chunked]
        
        if answer.strip():

            index.upsert(vectors=zip(doc_ids, docs_embed, metadatas))

        return jsonify(success=True)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error updating answer"), 500
    
@app.route("/edit-answer", methods=["POST"])
def edit_answer():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    question_id = data["question_id"]
    question = data["question"]
    answer = data["answer"]

    try:
        cur.execute(
            "UPDATE interview_questions SET answer = %s WHERE id = %s",
            (answer, question_id),
        )
        db.commit()
        cur.close()

        embeddings = OpenAIEmbeddings(
            openai_api_key=openai_api_key,
            model="text-embedding-ada-002"
        )
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=5000,
            chunk_overlap=20,
            length_function=len,
        )
        docs_chunked = text_splitter.create_documents([f"Question: {question}, Answer: {answer}"])
        docs_text = [doc.page_content for doc in docs_chunked]
        docs_embed = embeddings.embed_documents(docs_text)

        if answer.strip():

            index.update(
                id=f"{question_id}",
                values=docs_embed,
                namespace=""
            )

        return jsonify(success=True)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error updating answer"), 500

@app.route("/generate-recommendations", methods=["POST"])
def generate_recommendations():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    job_id = data.get("job_id")
    version_id = data.get("version_id")

    chat = ChatOpenAI(
        model_name="gpt-3.5-turbo",
        openai_api_key=openai_api_key,
        temperature=0,
    )

    cur.execute(f"SELECT * FROM saved_jobs WHERE id = %s AND user_id = %s", (job_id, user_id,))

    job = cur.fetchall()

    cur.execute(f"SELECT * FROM resume WHERE section = %s AND user_id = %s", ("FULL RESUME",user_id,))
    
    resume = cur.fetchone()[3]

    cur.execute(f"SELECT * FROM resume_recommendations WHERE user_id = %s AND job_id = %s", (user_id, job_id,))

    existing_recommendations = cur.fetchall()
    recommendations_list = [row[2] for row in existing_recommendations]
    recommendations_str = ', '.join(f'"{question}"' for question in recommendations_list)

    embeddings = OpenAIEmbeddings(
        openai_api_key=openai_api_key,
        model="text-embedding-ada-002"
    )

    query_job = embeddings.embed_query("What are the important responsibilities, qualifications, skills, and requirements for this job?")

    jobs_docs = index.query(
        vector=query_job,
        filter={
            "type": {"$eq": "jobs"},
            "job_id": str(job_id),
            "user": f"""{user_id}"""
        },
        top_k=4,
        include_metadata=True
    )

    query_resume = embeddings.embed_query("What are the important responsibilities, qualifications, skills, and requirements outlined in this resume?")

    resume_docs = index.query(
        vector=query_resume,
        filter={
            "type": {"$eq": "resume"},
            "user": f"""{user_id}"""
        },
        top_k=2,
        include_metadata=True,
        namespace=""
    )

    jobs_docs_texts = jobs_docs["matches"]

    jobs_docs_texts_flat = [doc["metadata"]["text"] for doc in jobs_docs_texts]

    jobs_docs_str = "\n\n".join(jobs_docs_texts_flat)

    resume_docs_texts = resume_docs["matches"]

    resume_docs_texts_flat = [doc["metadata"]["text"] for doc in resume_docs_texts]

    resume_docs_str = "\n\n".join(resume_docs_texts_flat)

    template_questions_base = f"""Based on the following details about my resume and the job I am applying to, please provide me 5 recommendations on how I can update my resume to increase my chances of landing an interview. The recommendations should cover different areas such as skills, work experience, projects, and format. Each recommendation should be justified in terms of how it matches with the job description or enhances my profile. Please return the recommendations separated by newlines.
        
    Job Description Information: {jobs_docs_str}"""

    if existing_recommendations:
        exclude_similar = f"""These recommendations have already been provided: [{recommendations_str}]"""
    else:
        exclude_similar = ""

    template_final = template_questions_base + """\nResume: {context}\n """ + exclude_similar


    chain = LLMChain(
        llm=chat,
        verbose=True,
        prompt=PromptTemplate.from_template(template_final)
    )

    try:
        recommendation_response = chain(resume)
        print(recommendation_response)
        # Convert the questions string into a Python list
        recommendation_string = recommendation_response['text']
        recommendation_list = [re.sub(r'^\d+\.\s*', '', i) for i in recommendation_string.split('\n')]

        # Save the questions to the interview_questions table
        for recommendation in recommendation_list:
            cur.execute(
                "INSERT INTO resume_recommendations (user_id, job_id, recommendation) VALUES (%s, %s, %s)",
                (user_id, job_id, recommendation),
            )

        # Commit the changes to the database
        db.commit()
        cur.close()

        return jsonify(recommendations=recommendation_list)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error fetching GPT-3.5 API"), 500

@app.route("/get-recommendations", methods=["POST"])
def get_recommendations():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    version_id = data.get("version_id")
    job_id = data.get("job_id")

    print(data)

    # Fetch recommendations based on user_id and version_id
    # Replace 'recommendations' with the appropriate table name and columns
        
    cur.execute("SELECT * FROM resume_recommendations WHERE user_id = %s AND job_id = %s", (user_id, job_id))
        
    recommendations = cur.fetchall()
    db.commit()
    cur.close()

    print("Recommendations fetched:", recommendations) 

    recommendations_data = [
        {
            "id": rec[0],
            "user_id": rec[1],
            "version_id": rec[2],
            "recommendation" : rec[4]

        }
        for rec in recommendations
    ]

    return jsonify(recommendations=recommendations_data)


@app.route("/delete-recommendation", methods=["POST"])
def delete_recommendation():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    recommendation_id = data["recommendation_id"]

    try:
        # Delete the recommendation based on recommendation_id and user_id
        # Replace 'recommendations' with the appropriate table name
        cur.execute("DELETE FROM resume_recommendations WHERE id = %s AND user_id = %s", (recommendation_id, user_id))
        db.commit()
        cur.close()
        return jsonify(success=True, message="Recommendation deleted successfully")
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error deleting the recommendation")


@app.route("/get-jobs", methods=["POST"])
def get_jobs():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    page = int(data.get("page", 1))
    items_per_page = int(data.get("itemsPerPage", 10))

    cur.execute("SELECT * FROM saved_jobs WHERE user_id = %s ORDER BY date_created DESC LIMIT %s OFFSET %s", (user_id, items_per_page, (page - 1) * items_per_page))
    saved_jobs = cur.fetchall()
    cur.close()

    jobs = [
        {
            "id": job[0],
            "user_id": job[1],
            "job_title": job[2],
            "company_name": job[3],
            "job_description": job[4],
            "status": job[5],
            "post_url": job[6],
            "date_added": job[7],
        }
        for job in saved_jobs
    ]

    return jsonify(jobs=jobs, headers=["Job Title/Role", "Company Name", "Job Description", "Status"])

@app.route("/search-jobs", methods=["POST"])
def search_jobs():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    page = int(data.get("page", 1))
    items_per_page = int(data.get("itemsPerPage", 10))
    query = data.get("query")
    location = data.get("location")

    threshold = datetime.now() - timedelta(days=1)

    # check if there is a row that has a date_added timestamp within the last 24 hours
    cur.execute(
        "SELECT * FROM jobs WHERE date_added > %s AND user_id = %s",
        (threshold,
        user_id),
    )
    recent_jobs = cur.fetchall()
    cur.execute("SELECT job_title FROM users WHERE id = %s", (user_id,))
    user_job_title = cur.fetchone()

    if not recent_jobs:
        
        if not user_job_title:

            embeddings = OpenAIEmbeddings(
                openai_api_key=openai_api_key,
                model="text-embedding-ada-002"
            )

            query_embed = embeddings.embed_query("Current job title and current location/address")

            initialize_profile_docs = index.query(
                vector=query_embed,
                filter={
                    "type": {"$eq": "resume"},
                    "user": f"""{user_id}"""
                },
                top_k=2,
                include_metadata=True,
                namespace=""
            )

            init_docs = initialize_profile_docs["matches"]
            init_docs_texts_flat = [doc["metadata"]["text"] for doc in init_docs]
            init_docs_str = "\n\n".join(init_docs_texts_flat)

            chat = ChatOpenAI(
                model_name="gpt-3.5-turbo",
                openai_api_key=openai_api_key,
                temperature=0,
            )

            template_final = """\nBelow is information from a person's resume. From this information, determine the user's current job title. Return only the job title and nothing else.\n  Here is the information from their resume: {context}"""

            chain = LLMChain(
                llm=chat,
                prompt=PromptTemplate.from_template(template_final)
            )

            initialize_user = chain(init_docs_str)

            user_job_title = initialize_user['text']
    
        # Call the SERP API job search function if no recent jobs in the database
        job_results = call_serp_api(user_job_title, location)
        # Delete all rows from the jobs table
        cur.execute("DELETE FROM jobs WHERE user_id = %s", (user_id,))
        # Save the job results to the database
        print(job_results)
        for job in job_results['jobs_results']:
            cur.execute("""
                INSERT INTO jobs (user_id, title, company_name, location, description, job_highlights, source, extensions, date_added)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, current_timestamp)
            """, (user_id, job["title"], job["company_name"], job["location"], job["description"], ', '.join(job["job_highlights"][0]['items']), job["via"],', '.join(job["extensions"])))
        db.commit()

    # Fetch jobs from the database
    cur.execute("SELECT * FROM jobs WHERE user_id = %s ORDER BY date_added DESC LIMIT %s OFFSET %s", (user_id, items_per_page, (page - 1) * items_per_page))
    jobs = cur.fetchall()
    cur.close()

    jobs = [
        {
            "id": job[0],
            "title": job[2],
            "company_name": job[3],
            "location": job[4],
            "description": job[5],
            "job_highlights": job[6],
            "date_added": job[7],
            "saved": job[9],
        }
        for job in jobs
    ]

    return jsonify(jobs=jobs, headers=["Job Title", "Company Name", "Location", "Job Description", "Job Highlights"])

@app.route("/create-job", methods=["POST"])
def create_job():
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    job_title = data["job_title"]
    company_name = data["company_name"]
    job_description = data["job_description"]
    status = data["status"]
    post_url = data.get("post_url") or ""
    saved = data.get("saved")
    id = data.get("id")
    date_created = datetime.utcnow()

    print(user_id, id, saved)

    cur.execute(
        "INSERT INTO saved_jobs (user_id, job_title, company_name, job_description, status, post_url, date_created) VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id",
        (user_id, job_title, company_name, job_description, status, post_url, date_created),
    )

    job_id = cur.fetchone()[0]

    if saved:
        cur.execute("UPDATE jobs SET saved = True WHERE user_id = %s AND id = %s", (user_id, id),)

    cur.execute(f"SELECT * FROM resume WHERE section = %s AND user_id = %s", ("FULL RESUME",user_id,))
    
    resume = cur.fetchone()[3]

    cur.execute(
        "INSERT INTO resume_versions (user_id, job_id, version_text) VALUES (%s, %s, %s) RETURNING id",
        (user_id, job_id, resume),
    )

    db.commit()

    job_string = f"""
    Job Title: {job_title}

    Company Name: {company_name}

    Job Description: {job_description}

    Status: {status}

    Post URL: {post_url}

    Date Created: {date_created}
    """

    try:

        db.commit()
        cur.close()

        embeddings = OpenAIEmbeddings(
            openai_api_key=openai_api_key,
            model="text-embedding-ada-002"
        )

        text_splitter = RecursiveCharacterTextSplitter(
            # Set a really small chunk size, just to show.
            chunk_size = 500,
            chunk_overlap  = 20,
            length_function = len,
        )

        docs_chunked = text_splitter.create_documents([job_string])
        docs_text = [doc.page_content for doc in docs_chunked]
        docs_embed = embeddings.embed_documents(docs_text)

        doc_ids = [str(uuid.uuid4()) for _ in docs_chunked]
        metadatas = [{"type": "jobs", "user": str(user_id), "text": _.page_content, "job_id": str(job_id), "company_name": company_name, "job_title": job_title} for _ in docs_chunked]

        index.upsert(vectors=zip(doc_ids, docs_embed, metadatas))

        return jsonify(
            success=True,
            id=job_id,
            user_id=user_id,
            job_title=job_title,
            company_name=company_name,
            job_description=job_description,
            status=status,
            post_url=post_url,
            date_created=date_created
        )

    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error updating answer"), 500


@app.route("/edit-job/<int:job_id>", methods=["PUT"])
def edit_joblication(job_id):
    db = get_db()
    cur = db.cursor()

    data = request.json
    user_id = data["user_id"]
    job_title = data["job_title"]
    company_name = data["company_name"]
    job_description = data["job_description"]
    status = data["status"]
    post_url = data["post_url"]

    cur.execute(
        "UPDATE saved_jobs SET user_id = %s, job_title = %s, company_name = %s, job_description = %s, status = %s, post_url = %s WHERE id = %s",
        (user_id, job_title, company_name, job_description, status, post_url, job_id),
    )
    db.commit()
    cur.close()

    return jsonify(success=True, job_id=job_id)

@app.route("/delete-job/<int:job_id>", methods=["DELETE"])
def delete_joblication(job_id):
    # Retrieve the user_id from the request body
    user_id = request.json.get('user_id')

    db = get_db()

    cur = db.cursor()
    # Update the SQL DELETE statement to delete the record that matches both user_id and job_id
    cur.execute(
        "DELETE FROM saved_jobs WHERE id = %s AND user_id = %s", (job_id, user_id)
    )
    db.commit()
    cur.close()

    index.delete(
        filter={
            "type": {"$eq": "jobs"},
            "job_id": job_id,
            "user": user_id
        }
    )

    return jsonify(success=True)

@app.route("/get-resume-versions", methods=["POST"])
def get_resume_versions():
    db = get_db()
    data = request.json
    user_id = data["user_id"]
    job_id = data.get("job_id")
    page = int(data.get("page", 1))
    items_per_page = int(data.get("itemsPerPage", 10))

    cur = db.cursor()

 
    cur.execute("SELECT * FROM resume_versions WHERE user_id = %s AND job_id = %s", (user_id, job_id,))
    
    version = cur.fetchone()

    resume_version = {
        "id": version[0],
        "user_id": version[1],
        "job_id": job_id,
        "version_name": version[3],
        "version_text": version[4]
    }

    return jsonify(resume_version)

@app.route("/edit-resume-version/<int:resume_version_id>", methods=["PUT"])
def edit_resume_version(resume_version_id):
    db = get_db()
    data = request.json
    user_id = data["user_id"]
    job_id = data["job_id"]
    version_name = data["version_name"]
    version_text = data["version_text"]

    cur = db.cursor()
    cur.execute(
        "UPDATE resume_versions SET user_id = %s, job_id = %s, version_name = %s, version_text = %s, WHERE id = %s",
        (user_id, job_id, version_name, version_text, resume_version_id),
    )
    db.commit()

    return jsonify(success=True, resume_version_id=resume_version_id)


@app.route("/delete-resume-version/<int:resume_version_id>", methods=["DELETE"])
def delete_resume_version(resume_version_id):
    db = get_db()

    cur = db.cursor()
    cur.execute(
        "DELETE FROM resume_versions WHERE id = %s", (resume_version_id,)
    )
    db.commit()

    return jsonify(success=True)

@app.route("/generate-cover-letter", methods=["POST"])
def generate_cover_letter():
    db = get_db()
    data = request.json
    user_id = data["user_id"]
    job_id = data.get("job_id")
    version_id = data.get("version_id")

    chat = ChatOpenAI(
        model_name="gpt-3.5-turbo",
        openai_api_key=openai_api_key,
        temperature=0,
    )

    cur = db.cursor()

    cur.execute(f"SELECT * FROM saved_jobs WHERE id = %s AND user_id = %s", (job_id, user_id,))

    job = cur.fetchone()

    cur.execute(f"SELECT * FROM resume WHERE section = %s AND user_id = %s", ("FULL RESUME",user_id,))
    
    resume = cur.fetchone()

    print(job[4])
    print(resume[3])

    template_questions_base = f"""Below is a user's resume and the description of a job they are applying to. Please write a cover letter for their job application. Do not make anything up.\n
        
    Job Description Information: {job[4]} \n"""

    template_final = template_questions_base + """Resume: {context}\n """

    chain = LLMChain(
        llm=chat,
        verbose=True,
        prompt=PromptTemplate.from_template(template_final)
    )

    try:
        response = chain(resume[3])
        print(response)
        # Convert the questions string into a Python list
        string = response['text']

        cur.execute("INSERT INTO cover_letter (user_id, job_id, cover_letter) VALUES (%s, %s, %s)", (user_id, job_id , string),)
        cover_letter_id = cur.lastrowid  # Get the ID of the inserted row

        # Commit the changes to the database
        db.commit()

        cover_letter_data = {
            "id": cover_letter_id,
            "user_id": user_id,
            "job_id": job_id,
            "cover_letter" : string
        }

        return jsonify(cover_letter_data)
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error fetching GPT-3.5 API"), 500
    
@app.route("/get-cover-letter", methods=["POST"])
def get_cover_letter():
    db = get_db()
    data = request.json
    user_id = data["user_id"]
    job_id = data.get("job_id")

    print(data)

    cur = db.cursor()

    # Fetch recommendations based on user_id and version_id
    # Replace 'recommendations' with the appropriate table name and columns

    cur.execute("SELECT * FROM cover_letter WHERE user_id = %s AND job_id = %s", (user_id, job_id))
        
    cover_letter = cur.fetchone()

    # If cover_letter is None, return a specific message
    if cover_letter is None:
        return jsonify(message="No cover letter found for the provided user_id and job_id."), 404

    cover_letter_data = {
            "id": cover_letter[0],
            "user_id": cover_letter[1],
            "job_id": cover_letter[2],
            "cover_letter" : cover_letter[3]
        }
    
    return jsonify(cover_letter_data)

@app.route("/save-cover-letter", methods=["POST"])
def save_cover_letter():
    db = get_db()
    data = request.json
    user_id = data["user_id"]
    job_id = data.get("job_id")
    new_cover_letter = data.get("cover_letter")

    print(data)

    cur = db.cursor()

    # Check if cover letter already exists for the user and job
    cur.execute("SELECT * FROM cover_letter WHERE user_id = %s AND job_id = %s", (user_id, job_id))
    cover_letter = cur.fetchone()

    if cover_letter is None:
        # Insert new cover letter
        cur.execute(
            "INSERT INTO cover_letter (user_id, job_id, cover_letter) VALUES (%s, %s, %s)",
            (user_id, job_id, new_cover_letter)
        )
    else:
        # Update existing cover letter
        cur.execute(
            "UPDATE cover_letter SET cover_letter = %s WHERE user_id = %s AND job_id = %s",
            (new_cover_letter, user_id, job_id)
        )

    db.commit()

    return jsonify(success=True, message="Cover letter updated successfully."), 200

@app.route("/delete-cover-letter", methods=["POST"])
def delete_cover_letter():
    db = get_db()
    data = request.json
    user_id = data["user_id"]
    cover_letter_id = data["cover_letter_id"]

    cur = db.cursor()
    try:
        # Delete the recommendation based on recommendation_id and user_id
        # Replace 'recommendations' with the appropriate table name
        cur.execute("DELETE FROM cover_letter WHERE id = %s AND user_id = %s", (cover_letter_id, user_id))
        db.commit()
        return jsonify(success=True, message="Cover letter deleted successfully")
    except Exception as e:
        print(e)
        return jsonify(success=False, message="Error deleting the cover letter")

@app.route("/get-database-tables", methods=["GET"])
def get_database_tables():
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'")
        rows = cur.fetchall()
        table_names = list(map(lambda row: {"name": row[0]}, rows))
        return jsonify(tables=table_names)
    except Exception as e:
        print("Error fetching database tables:", e)
        return jsonify(error="An error occurred while fetching database tables."), 500


@app.route("/get-table-data/<string:table_name>", methods=["GET"])
def get_table_data_route(table_name):
    user_id = request.args.get("user_id")
    
    try:
        table_data = get_table_data(table_name, user_id)
        print(table_data)
        return table_data
    except Exception as e:
        print(f"Error fetching data for table {table_name}:", e)
        return jsonify(error="An error occurred while fetching table data."), 500
    
@app.route("/delete-row/<string:table_name>/<int:row_id>", methods=["DELETE"])
def delete_row_route(table_name, row_id):
    user_id = request.args.get("user_id")
    
    try:
        deleted_rows = delete_row(table_name, row_id, user_id)
        return jsonify(deleted_rows=deleted_rows)
    except Exception as e:
        print(f"Error deleting row {row_id} from table {table_name}:", e)
        return jsonify(error="An error occurred while deleting the row."), 500

    

@app.route("/get-messages", methods=["GET"])
def get_messages_route():
    user_id = request.args.get("user_id")
    chat_id = request.args.get("chat_id")

    try:
        messages = get_messages_by_chat_id(chat_id, user_id)
        return jsonify(messages)
    except Exception as e:
        print(f"Error fetching messages for chat_id {chat_id}:", e)
        return jsonify(error="An error occurred while fetching messages."), 500

def get_messages_by_chat_id(chat_id, user_id):
    db = get_db()

    cur = db.cursor()

    cur.execute("SELECT * FROM messages WHERE chat_id = %s AND user_id = %s", (chat_id, user_id))
    rows = cur.fetchall()
        
    column_names = [column[0] for column in cur.description]
        
    result = [dict(zip(column_names, row)) for row in rows]
        
    return result


if __name__ == "__main__":
    app.run(port=3001)